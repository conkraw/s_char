import streamlit as st
from docx import Document
from docx.shared import Pt
import os
import re
import requests

# Function to format diagnosis names
def format_diagnosis_name(diagnosis):
    diagnosis = diagnosis.replace('_', ' ')
    formatted_name = re.sub(r'(?<!^)(?=[A-Z])', ' ', diagnosis)
    formatted_name = formatted_name.title()
    return formatted_name

# Function to fetch files from GitHub repository
def fetch_files_from_github(folder_name):
    url = f"https://api.github.com/repos/conkraw/s_char/contents/{folder_name}"
    files = []
    
    try:
        # Send GET request to GitHub API to fetch contents of the folder
        response = requests.get(url)
        
        if response.status_code == 200:
            # Parse the response JSON
            response_files = response.json()
            for file in response_files:
                # Filter out only .docx files
                if file['name'].endswith('.docx'):
                    files.append(file['name'])
        else:
            st.error(f"Failed to fetch files: Status code {response.status_code}")
            st.write(response.text)  # Display the response content for debugging
    except requests.exceptions.RequestException as e:
        st.error(f"An error occurred while fetching files: {e}")
        st.write(str(e))  # Display the exception details for debugging
    
    return files

# Function to download and extract content from a document
def fetch_file_content(folder_name, file_name):
    url = f"https://raw.githubusercontent.com/conkraw/s_char/master/{folder_name}/{file_name}"
    try:
        # Download the document as raw content
        response = requests.get(url)
        
        if response.status_code == 200:
            # Save the content as a .docx file locally
            with open(file_name, "wb") as f:
                f.write(response.content)
            
            # Now read the content from the local file
            doc = Document(file_name)
            content = "\n".join([para.text for para in doc.paragraphs])
            os.remove(file_name)  # Remove the local file after reading content
            return doc  # Return the Document object, not just plain text
        else:
            st.error(f"Failed to fetch content: Status code {response.status_code}")
            return None
    except requests.exceptions.RequestException as e:
        st.error(f"An error occurred while fetching content: {e}")
        return None

# Function to create a Word document with specific font settings and single spacing
def create_word_doc(text):
    doc = Document()
    
    for line in text.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    
    output_path = "updated_note.docx"
    doc.save(output_path)
    return output_path

def combine_notes(assess_text, critical_care_reason, diagnoses, free_text_diag=None, free_text_plan=None, physical_exam_day=None, ros_file=None, critical_care_time=None):
    doc = Document()

    # Add the introductory statement at the top (italicized, Arial, font size 9)
    intro_paragraph = doc.add_paragraph()
    intro_run = intro_paragraph.add_run(
        "I personally examined the patient separately and discussed the case with the resident/physician assistant and with any services involved in a multidisciplinary fashion. I agree with the resident/physician's assistant documentation with any exceptions noted below:"
    )
    intro_run.italic = True
    intro_run.font.name = 'Arial'
    intro_run.font.size = Pt(9)
    intro_paragraph.paragraph_format.space_after = Pt(0)
    intro_paragraph.paragraph_format.space_before = Pt(0)

    # Add "OVERNIGHT EVENTS:" section
    overnight_paragraph = doc.add_paragraph()
    overnight_header_run = overnight_paragraph.add_run("OVERNIGHT EVENTS:")
    overnight_header_run.bold = True
    overnight_header_run.underline = True
    overnight_header_run.font.name = 'Arial'
    overnight_header_run.font.size = Pt(9)
    overnight_content_run = overnight_paragraph.add_run(" No acute events were noted overnight.")
    overnight_content_run.font.name = 'Arial'
    overnight_content_run.font.size = Pt(9)
    overnight_paragraph.paragraph_format.space_after = Pt(6)
    overnight_paragraph.paragraph_format.space_before = Pt(6)

    # Add "SUBJECTIVE" header and ROS content
    if ros_file != "None.docx":
        ros_paragraph = doc.add_paragraph()
        
        # Add SUBJECTIVE heading (bold, underline)
        ros_run = ros_paragraph.add_run("SUBJECTIVE: ")
        ros_run.bold = True
        ros_run.underline = True
        ros_run.font.name = 'Arial'
        ros_run.font.size = Pt(9)

        ros_paragraph.paragraph_format.space_after = Pt(0)
        ros_paragraph.paragraph_format.space_before = Pt(0)
        
        # Fetch the content of the selected ROS file
        ros_doc = fetch_file_content('ros', ros_file)

        if ros_doc:
            for para in ros_doc.paragraphs:
                new_paragraph = doc.add_paragraph()

                # Split the paragraph text by the target phrases and apply formatting to those specific phrases
                text = para.text
                text_chunks = []

                # Check and split for "OVERNIGHT EVENTS"
                if "OVERNIGHT EVENTS" in text:
                    text_chunks.extend(text.split("OVERNIGHT EVENTS"))
                    text_chunks.insert(1, "OVERNIGHT EVENTS")
                else:
                    text_chunks.append(text)

                # Now handle applying bold/underline to "OVERNIGHT EVENTS" and "SUBJECTIVE"
                formatted_text = []
                for chunk in text_chunks:
                    if chunk == "OVERNIGHT EVENTS":
                        # Apply bold and underline only to "OVERNIGHT EVENTS"
                        run = new_paragraph.add_run(chunk)
                        run.bold = True
                        run.underline = True
                    elif "SUBJECTIVE" in chunk:
                        # Apply bold and underline to "SUBJECTIVE"
                        run = new_paragraph.add_run(chunk)
                        run.bold = True
                        run.underline = True
                    else:
                        # For normal text, just add as-is
                        run = new_paragraph.add_run(chunk)
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)

                new_paragraph.paragraph_format.space_after = Pt(6)
                new_paragraph.paragraph_format.space_before = Pt(0)

    # Add Objective section if a physical exam day is selected
    if physical_exam_day:
        objective_paragraph = doc.add_paragraph()
        objective_run = objective_paragraph.add_run("OBJECTIVE:")
        objective_run.bold = True
        objective_run.underline = True
        objective_run.font.name = 'Arial'
        objective_run.font.size = Pt(9)
        objective_paragraph.paragraph_format.space_after = Pt(0)
        objective_paragraph.paragraph_format.space_before = Pt(0)

        # Fetch the content of the selected physical exam day
        physical_exam_doc = fetch_file_content('physicalexam', physical_exam_day)

        # Add the fetched content under the OBJECTIVE section
        if physical_exam_doc:
            for para in physical_exam_doc.paragraphs:
                new_paragraph = doc.add_paragraph(para.text)
                new_paragraph.paragraph_format.space_after = Pt(0)
                new_paragraph.paragraph_format.space_before = Pt(0)
                for run in new_paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)

    # Add Assessment section
    assessment_paragraph = doc.add_paragraph()
    assessment_run = assessment_paragraph.add_run("ASSESSMENT:")
    assessment_run.bold = True
    assessment_run.underline = True
    assessment_run.font.name = 'Arial'
    assessment_run.font.size = Pt(9)
    assessment_paragraph.paragraph_format.space_after = Pt(0)
    assessment_paragraph.paragraph_format.space_before = Pt(6)

    assessment_content = doc.add_paragraph(assess_text)
    for run in assessment_content.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    # Add the "Why Critical Care" dropdown selection after assessment only if it's not empty
    if critical_care_reason != "":
        critical_care_paragraph = doc.add_paragraph()
        critical_care_run = critical_care_paragraph.add_run("CLINICAL INDICATIONS FOR CRITICAL CARE SERVICES:")
        critical_care_run.bold = True
        critical_care_run.underline = True
        critical_care_run.font.name = 'Arial'
        critical_care_run.font.size = Pt(9)
        critical_care_paragraph.paragraph_format.space_after = Pt(0)
        critical_care_paragraph.paragraph_format.space_before = Pt(0)

        critical_care_content = doc.add_paragraph(critical_care_reason)
        for run in critical_care_content.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)

    # Plan section
    plan_paragraph = doc.add_paragraph()
    plan_run = plan_paragraph.add_run("PLAN:")
    plan_run.bold = True
    plan_run.underline = True
    plan_run.font.name = 'Arial'
    plan_run.font.size = Pt(9)
    plan_paragraph.paragraph_format.space_after = Pt(0)
    plan_paragraph.paragraph_format.space_before = Pt(0)

    # Add selected diagnoses
    for i, diagnosis in enumerate(diagnoses, start=1):
        diagnosis_key = diagnosis.lower().replace(' ', '_') + '.docx'
        if os.path.exists(diagnosis_key):
            diagnosis_paragraph = doc.add_paragraph()
            diagnosis_run = diagnosis_paragraph.add_run(f"{i}). {diagnosis}")
            diagnosis_run.font.size = Pt(9)
            diagnosis_run.font.name = 'Arial'
            diagnosis_paragraph.paragraph_format.space_before = Pt(0)
            diagnosis_paragraph.paragraph_format.space_after = Pt(0)

            diagnosis_doc = Document(diagnosis_key)
            for para in diagnosis_doc.paragraphs:
                new_paragraph = doc.add_paragraph(para.text)
                new_paragraph.paragraph_format.space_before = Pt(0)
                new_paragraph.paragraph_format.space_after = Pt(0)

                for run in new_paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)

    # Append free-text diagnosis and plan if provided
    if free_text_diag and free_text_plan:
        doc.add_paragraph()  # Add a blank line
        doc.add_paragraph(f"Free Text Diagnosis: {free_text_diag}")
        doc.add_paragraph(f"Plan: {free_text_plan}")

    # Append Critical Care Time if provided
    if critical_care_time:
        plan_paragraph = doc.add_paragraph()
    
        # Add the bolded "Critical Care Time: " part
        plan_run = plan_paragraph.add_run("Critical Care Time: ")
        plan_run.bold = True
        plan_run.font.name = 'Arial'
        plan_run.font.size = Pt(9)
    
        # Add the non-bolded critical care time value
        plan_run = plan_paragraph.add_run(f"{critical_care_time}")
        plan_run.bold = False  # Remove bold formatting
        plan_run.font.name = 'Arial'
        plan_run.font.size = Pt(9)
    
        plan_paragraph.paragraph_format.space_before = Pt(6)
        plan_paragraph.paragraph_format.space_after = Pt(6)
    
    output_path = "combined_note.docx"
    doc.save(output_path)
    return output_path

# Title of the app
st.title("Note Management App")

# Header for the New Note section
st.header("Create a New Note")

# Input for room number
room_number = st.text_input("Enter Room Number:")

# Fetch both the diagnoses and physical exam days from GitHub
#physical_exam_days = fetch_files_from_github('physicalexam')
#ros_files = fetch_files_from_github('ros')

ros_files = {
    "None": "https://raw.githubusercontent.com/conkraw/s_char/main/ros/None.docx",
    "ROS_PARENT": "https://raw.githubusercontent.com/conkraw/s_char/main/ros/ros_parent.docx",
    "ROS_RN": "https://raw.githubusercontent.com/conkraw/s_char/main/ros/ros_rn.docx"
}

physical_exam_days = {
    "Adolescent Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day0.docx",
    "Infant Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day0.docx",
    "Child Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day0.docx",
    "Chronic Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day0.docx",
    
    "Adolescent Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day1.docx",
    "Infant Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day1.docx",
    "Child Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day1.docx",
    "Chronic Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day1.docx",
    
    "Adolescent Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day2.docx",
    "Infant Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day2.docx",
    "Child Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day2.docx",
    "Chronic Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day2.docx",
    
    "Adolescent Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day3.docx",
    "Infant Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day3.docx",
    "Child Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day3.docx",
    "Chronic Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day3.docx",
    
    "Adolescent Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day4.docx",
    "Infant Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day4.docx",
    "Child Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day4.docx",
    "Chronic Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day4.docx",
    
    "Adolescent Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day5.docx",
    "Infant Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day5.docx",
    "Child Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day5.docx",
    "Chronic Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day5.docx",
    
    "Adolescent Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day6.docx",
    "Infant Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day6.docx",
    "Child Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day6.docx",
    "Chronic Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day6.docx",
}


# Dynamically list available diagnosis documents in the current directory
available_docs = [f[:-5] for f in os.listdir('.') if f.endswith('.docx')]
formatted_conditions = [format_diagnosis_name(doc) for doc in available_docs]

# Sort the formatted conditions alphabetically
sorted_conditions = sorted(formatted_conditions)

# Add the selection input for physical exam day
if physical_exam_days:
    selected_exam_day = st.selectbox("Select Physical Examination Day:", physical_exam_days)
else:
    selected_exam_day = None

# Add the selection input for ROS file
if ros_files:
    selected_ros_file = st.selectbox("Select Review of Systems File:", ros_files)
else:
    selected_ros_file = None

# Select diagnoses
selected_conditions = st.multiselect("Choose diagnoses:", sorted_conditions)

assessment_text = st.text_area("Enter Assessment:")

critical_care_options = ["",
    "The patient requires critical care services due to the continuous management of invasive respiratory as well as hemodynamic support, which if not provided, would be life threatening to the patient.",
    "The patient requires critical care services due to the high risk of neurologic decompensation which could result in airway loss and respiratory failure, which is life threatening to the patient.",
    "The patient requires critical care services for management of the patient's airway and invasive mechanical respiratory support without which would be life threatening to the patient.",
    "The patient requires critical care services for management of the patient's airway and non-invasive mechanical respiratory support without which would be life threatening to the patient.",
    "The patient requires critical care services as the patient is at high risk of withdrawal, and thus requires intensive care monitoring.",
    "The patient is in a critically unstable condition with immediate risk of multisystem organ failure, including the cardiovascular and respiratory systems.",
    "The patient exhibits signs of rapid deterioration in respiratory and renal function, requiring intensive monitoring and therapeutic intervention.",
    "The patient is at high risk for cardiovascular collapse and respiratory failure, necessitating critical care for immediate stabilization.",
    "The patient's clinical presentation is consistent with impending respiratory and renal failure, with a high risk of further deterioration without intensive monitoring and intervention.",
    "The patient’s cardiovascular instability and compromised respiratory function place them at risk for immediate life-threatening complications.",
    "Signs of acute organ dysfunction, particularly in the cardiovascular and respiratory systems, suggest imminent need for intensive care to prevent further deterioration.",
    "The patient demonstrates significant hemodynamic instability, with potential for rapid deterioration in vital organ function, including the heart and lungs.",
    "Critical respiratory distress and compromised renal function require close monitoring and intervention to prevent irreversible organ damage.",
    "The patient is at risk of progressive multi-organ failure, particularly affecting the liver, kidneys, and lungs, and requires continuous intensive care management.",
    "The patient is experiencing severe cardiovascular instability and respiratory compromise, placing them at imminent risk for arrest without urgent critical care support.",
    "Acute exacerbation of chronic respiratory failure and unstable hemodynamics place the patient at significant risk for further deterioration, warranting critical care services.",
    "The patient’s rapidly worsening metabolic acidosis, along with cardiovascular and renal instability, indicates a need for continuous critical care to prevent life-threatening complications.",
    "Impaired cardiac output and acute respiratory distress syndrome (ARDS) necessitate intensive monitoring and immediate intervention to prevent further organ failure.",
    "The patient's clinical status is characterized by rapidly declining respiratory function, renal failure, and unstable blood pressure, with imminent risk for multi-organ dysfunction.",
    "There is significant concern for cardiovascular collapse and respiratory failure due to acute myocardial dysfunction, requiring intensive care interventions to stabilize the patient.",
    "The patient is critically ill with evidence of acute respiratory and renal failure, requiring immediate and sustained intensive care to support vital organ function.",
    "The patient is at risk of rapid deterioration in cardiac, respiratory, and renal function, necessitating continuous monitoring and aggressive therapeutic interventions.",
    "The patient's clinical presentation suggests an impending multi-organ system collapse, particularly affecting the heart, lungs, and kidneys, requiring intensive care to prevent irreversible damage."
]

    
selected_critical_care = st.selectbox("Why Critical Care:", critical_care_options)

# Add the Critical Care Time input (optional)
critical_care_time = st.text_input("Enter Critical Care Time (optional):")

if st.button("Submit New Note"):
    if selected_conditions and assessment_text and room_number:
        combined_file = combine_notes(
            assessment_text,
            selected_critical_care,
            selected_conditions,
            physical_exam_day=selected_exam_day,
            ros_file=selected_ros_file,
            critical_care_time=critical_care_time
        )
        file_name = f"{room_number}.docx"
        with open(combined_file, "rb") as f:
            st.download_button("Download Combined Note", f, file_name=file_name)
    else:
        st.error("Please fill out all fields.")



