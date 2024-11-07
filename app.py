import streamlit as st
from docx import Document
from docx.shared import Pt
import requests
from io import BytesIO

# Function to read the content of a .docx file from a URL
def read_docx_from_url(url):
    response = requests.get(url)
    doc = Document(BytesIO(response.content))
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return '\n'.join(content)

# Function to create a Word document with specific font settings and single spacing
# Function to create a Word document with specific font settings and single spacing
def create_word_doc(text, ros_text, physical_exam_text):
    doc = Document()

    # The specific phrase to add at the top of the document
    intro_text = (
        "I personally examined the patient separately and discussed the case with the resident/physician assistant "
        "and with any services involved in a multidisciplinary fashion. I agree with the resident/physician's assistant "
        "documentation with any exceptions noted below:"
    )
    
    # Add the introductory text as a paragraph, italicized, with Arial font and size 9
    intro_paragraph = doc.add_paragraph()
    intro_run = intro_paragraph.add_run(intro_text)
    intro_run.italic = True  # Set the intro text to be italicized
    intro_run.font.name = 'Arial'  # Set the font to Arial
    intro_run.font.size = Pt(9)   # Set the font size to 9
    
    # Add the "OVERNIGHT EVENTS:" header and its content
    overnight_paragraph = doc.add_paragraph()
    
    # "OVERNIGHT EVENTS:" header
    overnight_header_run = overnight_paragraph.add_run("OVERNIGHT EVENTS:")
    overnight_header_run.bold = True
    overnight_header_run.underline = True
    overnight_header_run.font.name = 'Arial'
    overnight_header_run.font.size = Pt(9)
    
    # "No acute events were noted overnight" content (no bold/underline)
    overnight_content_run = overnight_paragraph.add_run(" No acute events were noted overnight")
    overnight_content_run.font.name = 'Arial'
    overnight_content_run.font.size = Pt(9)
    
    # Add ROS if selected
    if ros_text:
        ros_paragraph = doc.add_paragraph()
        
        # Add ROS heading (e.g., "SUBJECTIVE:")
        ros_run = ros_paragraph.add_run("SUBJECTIVE: ")
        ros_run.bold = True
        ros_run.underline = True
        ros_run.font.name = 'Arial'
        ros_run.font.size = Pt(9)

        # Add the content of the ROS text
        ros_content_run = ros_paragraph.add_run("\n" + ros_text)
        ros_content_run.font.name = 'Arial'
        ros_content_run.font.size = Pt(9)
        
    # Now we handle the "OBJECTIVE:" section and the rest of the physical exam content
    physical_exam_lines = physical_exam_text.split("\n")
    
    # Add "OBJECTIVE:" directly to the same paragraph as the content
    physical_exam_paragraph = doc.add_paragraph()
    objective_run = physical_exam_paragraph.add_run("OBJECTIVE: ")  # Adding OBJECTIVE: header
    objective_run.bold = True
    objective_run.underline = True
    objective_run.font.name = 'Arial'
    objective_run.font.size = Pt(9)

    # Add the rest of the physical exam content in the same paragraph
    for line in physical_exam_lines:
        run = physical_exam_paragraph.add_run("\n" + line)
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    # Process the rest of the text passed into the function
    sections = text.split('\n')
    for section in sections:
        p = doc.add_paragraph()
        run = p.add_run(section)
        
        # Set font properties for the rest of the document
        run.font.name = 'Arial'
        run.font.size = Pt(9)

        # Apply bold and underline for specific sections
        if section.startswith("ASSESSMENT:"):
            run.bold = True
            run.underline = True
        elif section.startswith("PLAN:"):
            run.bold = True
            run.underline = True
        elif section.startswith("SUBJECTIVE:"):
            run.bold = True
            run.underline = True
        elif section.startswith("OBJECTIVE:"):
            run.bold = True
            run.underline = True
        elif section.startswith("OVERNIGHT EVENTS:"):
            run.bold = True
            run.underline = True
        elif section.startswith("CLINICAL INDICATIONS FOR CRITICAL CARE SERVICES:"):
            run.bold = True
            run.underline = True
            
        # Set single spacing
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)

    # Saving the final document with the required formatting applied
    output_path = "updated_note.docx"
    doc.save(output_path)
    return output_path


# Title of the app
st.title("Note Management App")

st.header("Update an Existing Note")

room_input = st.text_input("Enter Room Number:")

# Use session state to manage the text area content
if 'paragraph_text' not in st.session_state:
    st.session_state.paragraph_text = ""

st.session_state.paragraph_text = st.text_area("Enter the text for the note you want to update:", value=st.session_state.paragraph_text)

# Hardcode the URLs for ROS and Physical Exam files on GitHub
ros_files = {
    "None": "https://raw.githubusercontent.com/conkraw/s_char/main/ros/None.docx",
    "ROS_PARENT": "https://raw.githubusercontent.com/conkraw/s_char/main/ros/ros_parent.docx",
    "ROS_RN": "https://raw.githubusercontent.com/conkraw/s_char/main/ros/ros_rn.docx"
}

physical_exam_files = {
    "Adolescent Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day0.docx",
    "Infant Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day0.docx",
    "Child Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day0.docx",
    "Chronic Day 0": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day0.docx",
    
    "Adolescent Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day1.docx",
    "Infant Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day1.docx",
    "Child Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day1.docx",
    "Chronic Day 1": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day1.docx",
    
    "Adolescent Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day2.docx",
    "Infant Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day2.docx",
    "Child Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day2.docx",
    "Chronic Day 2": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day2.docx",
    
    "Adolescent Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day3.docx",
    "Infant Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day3.docx",
    "Child Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day3.docx",
    "Chronic Day 3": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day3.docx",
    
    "Adolescent Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day4.docx",
    "Infant Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day4.docx",
    "Child Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day4.docx",
    "Chronic Day 4": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day4.docx",
    
    "Adolescent Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day5.docx",
    "Infant Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day5.docx",
    "Child Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day5.docx",
    "Chronic Day 5": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day5.docx",
    
    "Adolescent Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Adolescent_Physical_Exam_Day6.docx",
    "Infant Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Infant_Physical_Exam_Day6.docx",
    "Child Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Child_Physical_Exam_Day6.docx",
    "Chronic Day 6": "https://raw.githubusercontent.com/conkraw/s_char/main/physicalexam/Chronic_Physical_Exam_Day6.docx",
}


# Dropdowns for selecting ROS and Physical Exam files
ros_selection = st.selectbox("Select ROS file:", list(ros_files.keys()))

sorted_physical_exam_options = sorted(physical_exam_files.keys(), key=lambda x: (x.split()[0], int(x.split()[2])))

physical_exam_selection = st.selectbox("Select Physical Exam file:", list(physical_exam_files.keys()))

# Allow the user to input their text for replacement
options = ["Continue", "Will continue", "We will continue", "We shall continue"]
col1, col2 = st.columns(2)

with col1:
    selected_option = st.selectbox("Select a phrase to replace:", options)

with col2:
    replacement = st.selectbox("Select a replacement phrase:", options)

# Construct the URLs for the selected files
ros_url = ros_files[ros_selection]
physical_exam_url = physical_exam_files[physical_exam_selection]

ros_text = read_docx_from_url(ros_url)  # Fetch the content of ROS file
physical_exam_text = read_docx_from_url(physical_exam_url)  # Fetch the content of Physical Exam file

if st.button("Replace"):
    if st.session_state.paragraph_text:
        # Perform replacement
        updated_text = st.session_state.paragraph_text.replace(selected_option, replacement)
        
        # Create the Word document
        word_file = create_word_doc(updated_text, ros_text, physical_exam_text)
        
        # Ensure room input is valid for filename
        if room_input:
            file_name = f"u_{room_input}.docx"
        else:
            file_name = "updated_note.docx"

        with open(word_file, "rb") as f:
            st.download_button("Download Updated Note", f, file_name=file_name)

        # Clear the text area
        st.session_state.paragraph_text = ""  # Clear the text area
        st.success("Replacement done! Text area cleared.")
    else:
        st.error("Please enter some text to update.")


