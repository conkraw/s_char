import streamlit as st
from docx import Document
from docx.shared import Pt
import os
import re

# Function to format diagnosis names
def format_diagnosis_name(diagnosis):
    # Replace underscores with spaces
    diagnosis = diagnosis.replace('_', ' ')
    # Add spaces before capital letters (camel case)
    formatted_name = re.sub(r'(?<!^)(?=[A-Z])', ' ', diagnosis)
    # Capitalize each word
    formatted_name = formatted_name.title()
    return formatted_name

# Function to create a Word document with specific font settings and single spacing
def create_word_doc(text):
    doc = Document()
    
    for line in text.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)  # No space after paragraph
        p.paragraph_format.space_before = Pt(0)  # No space before paragraph
        p.paragraph_format.line_spacing = Pt(12)  # Single spacing

    output_path = "updated_note.docx"
    doc.save(output_path)
    return output_path

# Function to combine diagnosis documents with formatted input text
def combine_notes(assess_text, diagnoses):
    doc = Document()
    
    # Assessment section
    assessment_paragraph = doc.add_paragraph()
    assessment_run = assessment_paragraph.add_run("ASSESSMENT:")
    assessment_run.bold = True
    assessment_run.underline = True
    assessment_run.font.name = 'Arial'
    assessment_run.font.size = Pt(9)
    assessment_paragraph.paragraph_format.space_after = Pt(0)  # No space after ASSESSMENT heading
    assessment_paragraph.paragraph_format.space_before = Pt(0)
    
    # Add the assessment text with formatting
    assessment_content = doc.add_paragraph(assess_text)
    for run in assessment_content.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
    assessment_content.paragraph_format.space_after = Pt(0)
    assessment_content.paragraph_format.space_before = Pt(0)

    # Plan section
    plan_paragraph = doc.add_paragraph()
    plan_run = plan_paragraph.add_run("PLAN:")
    plan_run.bold = True
    plan_run.underline = True
    plan_run.font.name = 'Arial'
    plan_run.font.size = Pt(9)
    plan_paragraph.paragraph_format.space_after = Pt(0)  # No space after PLAN heading
    plan_paragraph.paragraph_format.space_before = Pt(0)  # No space before PLAN heading

    for i, diagnosis in enumerate(diagnoses, start=1):
        # Convert the diagnosis back to the filename format
        diagnosis_key = diagnosis.lower().replace(' ', '_') + '.docx'
        if os.path.exists(diagnosis_key):
            # Add the diagnosis header with enhanced formatting
            diagnosis_paragraph = doc.add_paragraph()
            diagnosis_run = diagnosis_paragraph.add_run(f"{i}). {formatted_name}")
            #diagnosis_run.bold = True  # Bold the diagnosis
            diagnosis_run.font.size = Pt(10)  # Set font size
            diagnosis_run.font.name = 'Arial'  # Set font type
            diagnosis_paragraph.paragraph_format.space_after = Pt(0)  # No space after diagnosis
            diagnosis_paragraph.paragraph_format.space_before = Pt(0)  # No space before diagnosis

            # Add the content from the diagnosis document
            diagnosis_doc = Document(diagnosis_key)
            for para in diagnosis_doc.paragraphs:
                new_paragraph = doc.add_paragraph(para.text)
                for run in new_paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                new_paragraph.paragraph_format.space_after = Pt(0)  # No space after diagnosis content
                new_paragraph.paragraph_format.space_before = Pt(0)  # No space before diagnosis content

    output_path = "combined_note.docx"
    doc.save(output_path)
    return output_path

# Title of the app
st.title("Note Management App")

# Header for the New Note section
st.header("Create a New Note")

# Input for room number
room_number = st.text_input("Enter Room Number:")

# Dynamically list available diagnosis documents in the current directory
available_docs = [f[:-5] for f in os.listdir('.') if f.endswith('.docx')]
formatted_conditions = [format_diagnosis_name(doc) for doc in available_docs]

# Create a mapping for original document names
diagnosis_mapping = {format_diagnosis_name(doc): doc for doc in available_docs}

selected_conditions = st.multiselect("Choose diagnoses:", formatted_conditions)

assessment_text = st.text_area("Enter Assessment:")

if st.button("Submit New Note"):
    if selected_conditions and assessment_text and room_number:
        # Map selected conditions back to original filenames
        selected_conditions_original = [diagnosis_mapping[cond] for cond in selected_conditions]
        combined_file = combine_notes(assessment_text, selected_conditions_original)
        # Use room number in the filename
        file_name = f"{room_number}.docx"
        with open(combined_file, "rb") as f:
            st.download_button("Download Combined Note", f, file_name=file_name)
    else:
        st.error("Please fill out all fields.")

